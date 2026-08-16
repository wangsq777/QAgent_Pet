#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate the updated QAgent Pet AI product portfolio as a DOCX.

The visual system intentionally follows the retained PDF: A4 portrait,
white background, blue section rules, restrained blue table headers, and
compact Chinese body copy. The generated DOCX is the editable source for the
final PDF under output/pdf/.
"""

from __future__ import annotations

import os
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT_DOCX = ROOT / "output" / "docx" / "AI产品作品集-王思勤.docx"

# Reference-PDF palette.
BLUE = "2E75B6"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "EAF3FA"
PALE_BLUE = "F5F9FC"
INK = "333333"
MUTED = "666666"
LIGHT_MUTED = "888888"
GRID = "C7D3DD"
WHITE = "FFFFFF"

LATIN_FONT = "Arial"
# Arial Unicode MS is present on the target macOS workspace and is reliably
# discovered by headless LibreOffice, unlike the TTC-only STHeiti family.
CJK_FONT = "Arial Unicode MS"

# A4 width with symmetric 0.884-inch margins leaves exactly 9360 DXA.
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_run_font(
    run,
    *,
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = LATIN_FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), LATIN_FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), CJK_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, *, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = GRID, size: int = 4) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "start", "bottom", "end", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths: Sequence[int]) -> None:
    if sum(widths) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Table widths must sum to {CONTENT_WIDTH_DXA}: {widths}")

    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def paragraph_border_bottom(paragraph, color: str = BLUE, size: int = 7, space: int = 3) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)


def shade_paragraph(paragraph, fill: str, border: str | None = None) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    if border:
        p_bdr = p_pr.find(qn("w:pBdr"))
        if p_bdr is None:
            p_bdr = OxmlElement("w:pBdr")
            p_pr.append(p_bdr)
        for edge in ("top", "start", "bottom", "end"):
            node = OxmlElement(f"w:{edge}")
            node.set(qn("w:val"), "single")
            node.set(qn("w:sz"), "5")
            node.set(qn("w:space"), "4")
            node.set(qn("w:color"), border)
            p_bdr.append(node)


def keep_with_next(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    node = OxmlElement("w:keepNext")
    p_pr.append(node)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:cantSplit")
    tr_pr.append(node)


def set_style_font(style, size: float, color: str = INK, bold: bool = False) -> None:
    style.font.name = LATIN_FONT
    style.font.size = Pt(size)
    style.font.color.rgb = rgb(color)
    style.font.bold = bold
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), LATIN_FONT)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), CJK_FONT)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Inches(0.884)
    section.right_margin = Inches(0.884)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.68)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, 10.5)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = styles["Title"]
    set_style_font(title, 30, DARK_BLUE, True)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(6)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = styles["Subtitle"]
    set_style_font(subtitle, 15, MUTED, False)
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(10)
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    h1 = styles["Heading 1"]
    set_style_font(h1, 19, BLUE, True)
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.keep_with_next = True

    h2 = styles["Heading 2"]
    set_style_font(h2, 12.5, DARK_BLUE, True)
    h2.paragraph_format.space_before = Pt(9)
    h2.paragraph_format.space_after = Pt(5)
    h2.paragraph_format.keep_with_next = True

    h3 = styles["Heading 3"]
    set_style_font(h3, 11.3, INK, True)
    h3.paragraph_format.space_before = Pt(7)
    h3.paragraph_format.space_after = Pt(3)
    h3.paragraph_format.keep_with_next = True

    for list_style_name in ("List Bullet", "List Number"):
        style = styles[list_style_name]
        set_style_font(style, 10.2)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.10

    if "Portfolio Small" not in styles:
        small = styles.add_style("Portfolio Small", WD_STYLE_TYPE.PARAGRAPH)
    else:
        small = styles["Portfolio Small"]
    set_style_font(small, 8.8, MUTED)
    small.paragraph_format.space_before = Pt(0)
    small.paragraph_format.space_after = Pt(4)
    small.paragraph_format.line_spacing = 1.05

    # Quiet page numbers match the retained minimalist PDF.
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("QAgent Pet  |  ")
    set_run_font(run, size=8.5, color=LIGHT_MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), LIGHT_MUTED)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "17")
    r_pr.append(color)
    r_pr.append(size)
    r.append(r_pr)
    text = OxmlElement("w:t")
    text.text = "1"
    r.append(text)
    field.append(r)
    p._p.append(field)

    core = doc.core_properties
    core.title = "QAgent Pet - AI产品作品集"
    core.subject = "个人AI电子宠物伴侣：产品设计、Agent架构与桌宠MVP"
    core.author = "王思勤"
    core.keywords = "AI产品经理, Agent, 桌宠, 长期记忆, 情绪理解, Electron"
    core.comments = "基于QAgent项目仓库2026-07-23状态更新"


def add_spacer(doc: Document, points: float) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(points)


def add_text(
    doc: Document,
    text: str,
    *,
    bold_prefix: str | None = None,
    size: float = 10.5,
    color: str = INK,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    before: float = 0,
    after: float = 6,
    line: float = 1.10,
    italic: bool = False,
    keep: bool = False,
) -> object:
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if bold_prefix and text.startswith(bold_prefix):
        first = p.add_run(bold_prefix)
        set_run_font(first, size=size, color=color, bold=True)
        rest = p.add_run(text[len(bold_prefix):])
        set_run_font(rest, size=size, color=color, italic=italic)
    else:
        run = p.add_run(text)
        set_run_font(run, size=size, color=color, italic=italic)
    if keep:
        keep_with_next(p)
    return p


def add_bullet(doc: Document, text: str, *, level: int = 0, size: float = 10.2, after: float = 4) -> object:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5 + 0.25 * level)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    run = p.add_run(text)
    set_run_font(run, size=size, color=INK)
    return p


def add_number(doc: Document, text: str, *, size: float = 10.2, after: float = 5) -> object:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    run = p.add_run(text)
    set_run_font(run, size=size, color=INK)
    return p


def add_section_title(doc: Document, title: str, kicker: str | None = None) -> None:
    if kicker:
        add_text(doc, kicker.upper(), size=8.8, color=LIGHT_MUTED, after=2, keep=True)
    p = doc.add_paragraph(title, style="Heading 1")
    paragraph_border_bottom(p, BLUE, 7, 3)


def add_subheading(doc: Document, title: str) -> None:
    doc.add_paragraph(title, style="Heading 2")


def add_callout(doc: Document, label: str, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(8)
    p.paragraph_format.right_indent = Pt(8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.08
    label_run = p.add_run(f"{label}  ")
    set_run_font(label_run, size=10.4, color=BLUE, bold=True)
    body = p.add_run(text)
    set_run_font(body, size=10.4, color=INK)
    shade_paragraph(p, PALE_BLUE, GRID)


def add_table(
    doc: Document,
    headers: Sequence[str],
    rows: Iterable[Sequence[str]],
    widths: Sequence[int],
    *,
    font_size: float = 9.2,
    header_font_size: float = 9.4,
    first_col_bold: bool = False,
    alternate_fill: bool = False,
) -> object:
    data = list(rows)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.style = "Table Grid"

    header_row = table.rows[0]
    repeat_table_header(header_row)
    prevent_row_split(header_row)
    for index, header in enumerate(headers):
        cell = header_row.cells[index]
        set_cell_shading(cell, BLUE)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(str(header))
        set_run_font(run, size=header_font_size, color=WHITE, bold=True)

    for row_index, values in enumerate(data):
        row = table.add_row()
        prevent_row_split(row)
        for index, value in enumerate(values):
            cell = row.cells[index]
            if alternate_fill and row_index % 2 == 1:
                set_cell_shading(cell, "F7FAFC")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.03
            run = p.add_run(str(value))
            set_run_font(run, size=font_size, color=INK, bold=first_col_bold and index == 0)

    set_table_geometry(table, widths)
    set_table_borders(table)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(4)
    return table


def add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_break(WD_BREAK.PAGE)


def build_cover(doc: Document) -> None:
    add_spacer(doc, 86)
    add_text(
        doc,
        "AI PRODUCT PORTFOLIO",
        size=9.5,
        color=BLUE,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=15,
    )
    doc.add_paragraph("QAgent Pet", style="Title")
    doc.add_paragraph("个人 AI 电子宠物伴侣", style="Subtitle")
    add_text(
        doc,
        "让 AI 从一次性工具，变成有记忆、有边界、有主动性的桌面伙伴",
        size=11.2,
        color=MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=28,
    )
    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(26)
    paragraph_border_bottom(rule, GRID, 5, 2)

    cover_meta = [
        ("项目属性", "个人项目 / AI 产品从 0 到 1"),
        ("核心角色", "AI 产品经理、原型开发与迭代负责人"),
        ("产品形态", "Web 宠物控制中心 + Electron 桌宠 MVP"),
        ("开发周期", "2026.04 - 2026.07（持续迭代）"),
        ("当前阶段", "核心闭环已落地，进入桌宠体验增强与用户验证"),
    ]
    for label, value in cover_meta:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(5)
        r1 = p.add_run(f"{label}：")
        set_run_font(r1, size=10.2, color=BLUE, bold=True)
        r2 = p.add_run(value)
        set_run_font(r2, size=10.2, color=INK)

    add_spacer(doc, 28)
    add_text(doc, "王思勤", size=12.5, color=DARK_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_text(
        doc,
        "源于腾讯 AI 产品校园大赛 QQ 赛道，现已演进为个人桌面端关系型 AI 产品",
        size=8.8,
        color=LIGHT_MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        italic=True,
    )


def build_toc(doc: Document) -> None:
    add_section_title(doc, "目录", "CONTENTS")
    entries = [
        "01  产品概述与阶段演进",
        "02  用户问题、定位与验证假设",
        "03  核心体验闭环与功能架构",
        "04  AI 产品设计：记忆与情绪",
        "05  桌宠 MVP：双窗口与隐私边界",
        "06  Agent 扩展：GitHub 陪学与宠物串门",
        "07  技术架构、选型与安全治理",
        "08  版本迭代与关键决策",
        "09  实现证据与个人贡献",
        "10  路线图、验证指标与复盘",
    ]
    for entry in entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(7)
        p.paragraph_format.left_indent = Pt(8)
        run = p.add_run(entry)
        set_run_font(run, size=11.4, color=INK)

    add_subheading(doc, "本次更新重点")
    add_bullet(doc, "产品定位从 QQ 赛道 Demo 更新为“个人 AI 电子宠物伴侣”，明确 Web 完整能力中心与桌宠轻入口的分工。")
    add_bullet(doc, "补入已落地的 Electron 桌宠、2 字低敏提醒气泡、勿扰与托盘、五维情感结构、GitHub 陪学。")
    add_bullet(doc, "所有成果基于本地仓库核验；待实现能力和待验证指标单独标注，不包装为现有成绩。")


def build_overview(doc: Document) -> None:
    add_section_title(doc, "01  产品概述与阶段演进", "PRODUCT OVERVIEW")
    add_text(
        doc,
        "QAgent Pet 是一个围绕长期关系构建的个人 AI 电子宠物系统。它不只回答问题，而是通过稳定人格、长期记忆、情绪理解、主动关怀、亲密度成长和桌面常驻，让用户感受到“这只宠物认识我，也会主动想起我”。",
        after=8,
    )
    add_callout(doc, "一句话定位", "一个住在桌面上的关系型 AI 伙伴：完整能力在 Web 控制中心，轻量陪伴发生在桌宠窗口。")

    add_subheading(doc, "从竞赛 Agent 到桌面陪伴产品")
    add_table(
        doc,
        ["阶段", "产品形态", "核心目标"],
        [
            ["起点", "QQ 赛道 AI 宠物 Demo", "验证多角色对话、情绪感知与主动关怀"],
            ["能力扩展", "Web 宠物控制中心", "补齐记忆、画像、自定义宠物、串门与陪学"],
            ["当前", "Electron 桌宠 MVP", "跑通桌面常驻、轻气泡、轻聊天与完整面板闭环"],
        ],
        [1500, 3000, 4860],
        first_col_bold=True,
        alternate_fill=True,
    )

    add_subheading(doc, "核心价值主张")
    add_bullet(doc, "人格连续：预置宠物与自定义宠物拥有独立人设、语气、口头禅和互动节奏。")
    add_bullet(doc, "记忆延续：最近对话、向量召回和长期摘要共同维持跨时间的关系连续性。")
    add_bullet(doc, "主动但克制：主动关怀有宠物差异，也有勿扰、频率和隐私边界。")
    add_bullet(doc, "关系可成长：亲密度、陪伴时长、连续互动和学习奖励把互动反馈为长期关系。")


def build_positioning(doc: Document) -> None:
    add_section_title(doc, "02  用户问题、定位与验证假设", "USER & POSITIONING")
    add_text(
        doc,
        "MVP 阶段不先假设“人人需要 AI 陪伴”，而是从三个可观察场景出发：用户想倾诉但不想打扰朋友；想获得轻提醒但拒绝监控式通知；想持续学习但缺少低压力的陪伴与反馈。",
        after=8,
    )

    add_subheading(doc, "目标场景与产品解法")
    add_table(
        doc,
        ["用户场景", "核心阻力", "QAgent Pet 解法"],
        [
            ["独处时想说话", "真人社交有负担，通用助手缺少关系感", "稳定人格 + 情绪理解 + 长期记忆"],
            ["工作学习容易分心", "计划工具强调管理，容易造成压力", "宠物陪学 + 章末反馈 + 轻奖励"],
            ["需要提醒又怕打扰", "完整通知暴露隐私，频繁弹窗造成反感", "2 字概括 + 点击展开 + 勿扰模式"],
            ["喜欢个性化角色", "固定角色难长期匹配偏好", "3 个预置人格 + AI 自定义宠物"],
        ],
        [2400, 3180, 3780],
        font_size=8.9,
        first_col_bold=True,
        alternate_fill=True,
    )

    add_subheading(doc, "差异化定位")
    add_table(
        doc,
        ["维度", "通用 AI 助手", "角色聊天产品", "QAgent Pet"],
        [
            ["主要预期", "效率与答案", "剧情与角色扮演", "关系连续与日常陪伴"],
            ["主动性", "任务触发为主", "多数被动回复", "按人格与状态轻量主动"],
            ["记忆", "上下文/账户记忆", "依产品而异", "短期 + 语义召回 + 长期摘要"],
            ["入口", "App/网页", "App/网页", "Web 完整中心 + 桌面常驻轻入口"],
        ],
        [1500, 2300, 2300, 3260],
        font_size=8.7,
        first_col_bold=True,
    )
    add_callout(doc, "验证原则", "以上是产品假设，不是用户研究结论。下一阶段需要用提醒点击率、7 日留存、记忆命中满意度和勿扰开启率验证真实价值。")


def build_core_loop(doc: Document) -> None:
    add_section_title(doc, "03  核心体验闭环与功能架构", "CORE EXPERIENCE")
    add_subheading(doc, "一条关系型 AI 的最小闭环")
    steps = [
        "选择宠物：从预置人格进入，或配置名称、类型、性格、口头禅和习惯生成专属宠物。",
        "自然互动：对话中识别情绪与需求，按宠物人格生成不同语气的回应。",
        "形成记忆：保存最近消息、召回相关历史，并在话题变化时沉淀长期摘要与用户画像。",
        "关系反馈：互动、情绪支持、主动关怀回应和学习完成共同影响亲密度与状态。",
        "再次触达：Web 控制中心承载完整能力；桌宠用 2 字气泡低打扰召回，点击后展开完整内容。",
    ]
    for step in steps:
        add_number(doc, step)

    add_subheading(doc, "功能架构")
    add_table(
        doc,
        ["层级", "已实现能力", "用户感知"],
        [
            ["关系层", "人格、亲密度、主动关怀、轻养成状态", "宠物有性格，关系会变化"],
            ["认知层", "短期记忆、向量检索、长期摘要、用户画像", "它记得过去，也更了解我"],
            ["互动层", "聊天、情绪支持、日程、天气、日常分享", "能倾诉，也能处理轻任务"],
            ["扩展层", "自定义宠物、宠物串门、GitHub 陪学", "角色可创造，陪伴可进入具体场景"],
            ["终端层", "Web 三栏控制中心、Electron 桌宠、托盘与轻聊天", "完整能力与低打扰入口分离"],
        ],
        [1550, 4300, 3510],
        font_size=8.8,
        first_col_bold=True,
        alternate_fill=True,
    )
    add_callout(doc, "设计约束", "所有功能都必须回答同一个问题：它是否让宠物更像“长期关系对象”，而不是把更多工具堆进聊天框。")


def build_ai_design(doc: Document) -> None:
    add_section_title(doc, "04  AI 产品设计：记忆与情绪", "AI PRODUCT DESIGN")
    add_subheading(doc, "设计决策一：分层记忆，而不是无限堆上下文")
    add_table(
        doc,
        ["记忆层", "机制", "解决的问题"],
        [
            ["短期上下文", "最近 10 条完整消息", "保证当前话题连贯"],
            ["相关历史", "Embedding + 余弦相似度召回", "找回不在最近窗口中的关键事实"],
            ["长期摘要", "话题变化 / 轮次兜底触发 LLM 压缩", "降低上下文成本，保留关系事件"],
            ["用户画像", "后台 Agent 提取地区、身份、兴趣等字段", "为后续关怀和个性化提供稳定信号"],
        ],
        [1900, 3900, 3560],
        font_size=8.8,
        first_col_bold=True,
    )

    add_subheading(doc, "设计决策二：五维情感信号只在系统内部使用")
    add_table(
        doc,
        ["字段", "含义", "产品用途"],
        [
            ["reply", "宠物最终回复", "前端只展示这一层"],
            ["emotion", "当前情绪类别", "选择回应策略、记录趋势"],
            ["need", "陪伴/倾诉/认可/建议等需求", "决定支持方式与亲密度加成"],
            ["intensity", "1-5 级强度", "区分轻微波动与高强度情绪"],
            ["risk_level", "none / low / high", "高风险时切换安全回应"],
        ],
        [1600, 3000, 4760],
        font_size=8.7,
        first_col_bold=True,
        alternate_fill=True,
    )
    add_bullet(doc, "不在前端给用户贴“你现在很焦虑”之类的标签，避免误判造成冒犯。", size=9.8)
    add_bullet(doc, "高风险信号优先生成安全回应，引导联系现实可信任的人；产品不冒充专业心理服务。", size=9.8)
    add_bullet(doc, "MoodAgent 和 UserProfileAgent 在后台运行，失败不阻塞主对话。", size=9.8)


def build_desktop(doc: Document) -> None:
    add_section_title(doc, "05  桌宠 MVP：双窗口与隐私边界", "DESKTOP PET MVP")
    add_text(
        doc,
        "产品从 Web 走向桌面时，没有把完整聊天页缩进一个小浮窗，而是拆成“桌宠轻入口 + Web 完整中心”。这让常驻陪伴保持轻量，同时保留复杂功能的可用空间。",
        after=8,
    )
    add_subheading(doc, "双窗口分工")
    add_table(
        doc,
        ["终端", "承担任务", "关键交互"],
        [
            ["桌宠小窗口", "陪伴感、轻触达、快速对话", "透明无边框、置顶、拖拽、2 字气泡、轻聊天"],
            ["Web 完整中心", "聊天、记忆、学习、串门、自定义与设置", "三栏 App Shell、状态侧栏、功能导航"],
            ["系统托盘", "后台驻留与全局控制", "显示桌宠、打开面板、勿扰、切换宠物、退出"],
        ],
        [1800, 3600, 3960],
        font_size=9.0,
        first_col_bold=True,
        alternate_fill=True,
    )

    add_subheading(doc, "为什么提醒气泡只有 2 个字")
    add_bullet(doc, "可扫读：用“想你”“提醒”“陪学”等概括告诉用户发生了什么。")
    add_bullet(doc, "保隐私：桌面上不直接暴露完整回复、日程或情绪内容。")
    add_bullet(doc, "降打扰：用户点击宠物或气泡后，再进入轻聊天查看完整内容。")

    add_subheading(doc, "已落地与明确边界")
    add_table(
        doc,
        ["已实现 MVP", "暂不读取 / 暂未实现"],
        [
            ["后端端口检测与自动拉起", "屏幕内容、窗口标题、聊天软件内容"],
            ["桌宠与 Web 共用 session 和记忆", "开机自启、位置记忆、多显示器适配"],
            ["勿扰、托盘、切换预置宠物", "正式签名安装包与完整数据迁移"],
        ],
        [4680, 4680],
        font_size=8.9,
    )


def build_agents(doc: Document) -> None:
    add_section_title(doc, "06  Agent 扩展：GitHub 陪学与宠物串门", "AGENT EXPERIENCES")
    add_subheading(doc, "GitHub 项目陪学：把陪伴带入具体任务")
    add_text(
        doc,
        "用户输入公开 GitHub 仓库链接后，系统分析目录与关键文件，生成 3-6 章学习大纲；源码老师负责结构化讲解，当前宠物负责章末旁白、鼓励与关系反馈。",
        after=6,
    )
    add_table(
        doc,
        ["角色", "职责边界", "价值"],
        [
            ["源码老师 Agent", "解释为什么这样设计、运行时如何流转；只讲当前章节", "保证专业性与教学结构"],
            ["宠物伙伴 Agent", "按自身人格给章末反馈、回答陪伴类提问", "把学习压力转化为陪伴体验"],
            ["业务编排层", "进度、问答路由、章节奖励、防重复发奖", "让 Agent 输出进入可控产品流程"],
        ],
        [2100, 4300, 2960],
        font_size=8.8,
        first_col_bold=True,
        alternate_fill=True,
    )
    add_bullet(doc, "每章完成亲密度 +2；全部完成额外 +5，奖励记录防止重复领取。", size=9.8)
    add_bullet(doc, "仓库 README、代码和注释都被视为数据，不允许覆盖系统指令。", size=9.8)

    add_subheading(doc, "宠物串门：多角色体验中的上下文隔离")
    add_table(
        doc,
        ["防护层", "做法", "目的"],
        [
            ["Prompt 结构", "发言宠物完整人格在 system；对方信息降级为受限 context", "防止对方人格污染"],
            ["显式规则", "要求只扮演自己，不模仿对方", "约束角色边界"],
            ["输出后处理", "检查并移除可能的对方名字前缀", "降低角色串台"],
            ["记忆沉淀", "串门结束生成摘要，分别写入双方长期记忆", "让一次互动影响后续关系"],
        ],
        [1700, 4900, 2760],
        font_size=8.7,
        first_col_bold=True,
    )


def build_architecture(doc: Document) -> None:
    add_section_title(doc, "07  技术架构、选型与安全治理", "TECHNICAL ARCHITECTURE")
    add_subheading(doc, "当前架构")
    add_table(
        doc,
        ["层", "实现", "产品考量"],
        [
            ["桌面层", "Electron 主进程、预加载桥接、桌宠/轻聊天渲染页", "复用 Web 资产，快速验证桌面闭环"],
            ["表现层", "原生 HTML / CSS / JavaScript", "MVP 迭代快、依赖轻"],
            ["服务层", "FastAPI + 30 个 API 端点", "异步等待 LLM / GitHub / 天气服务"],
            ["AI 层", "MiniMax Anthropic Messages 兼容接口 + 分层 Prompt", "支持多调用链路、保留供应商切换空间"],
            ["数据层", "SQLite + aiosqlite + WAL + 12 张表", "单机 MVP 零运维，便于本地数据连续性"],
            ["记忆层", "短期消息、长期摘要、Embedding 向量表", "在成本与记忆体验间折中"],
        ],
        [1500, 4400, 3460],
        font_size=8.5,
        first_col_bold=True,
        alternate_fill=True,
    )

    add_subheading(doc, "安全不是发布前检查，而是产品约束")
    add_table(
        doc,
        ["风险", "已实现治理", "仍需继续"],
        [
            ["身份与越权", "X-User-Id 校验、session / pet 归属检查、API Key 中间件", "生产环境强制配置 API Key"],
            ["Prompt 注入", "输入清洗、仓库内容隔离、工具白名单与参数校验", "迁移为平台级结构化 tool calling"],
            ["网络与资源", "GitHub SSRF 防护、禁重定向、请求体限制、限流", "全局 LLM 并发上限"],
            ["本地数据", "SQLite 权限尝试设为 0o600、WAL、索引", "应用数据目录、备份恢复、正式持久化"],
        ],
        [1900, 4500, 2960],
        font_size=8.4,
        first_col_bold=True,
    )
    add_text(doc, "技术边界说明：当前向量召回仍在 SQLite 候选集内用 Python 计算余弦相似度，适合 MVP，不包装成生产级向量基础设施。", size=8.8, color=MUTED, italic=True, after=0)


def build_timeline(doc: Document) -> None:
    add_section_title(doc, "08  版本迭代与关键决策", "ITERATION")
    add_text(doc, "迭代不是按技术模块堆功能，而是围绕“关系连续 - 场景扩展 - 桌面触达 - 安全收口”逐层推进。", after=8)
    add_table(
        doc,
        ["时间 / 版本", "阶段目标", "关键交付"],
        [
            ["2026.05 / v1.0", "核心关系闭环", "3 个预置人格、聊天、亲密度、主动关怀、日程与天气"],
            ["2026.05 / v1.1-v1.2", "记忆与角色一致性", "双通道记忆、8 字段画像、口头禅概率控制"],
            ["2026.06 / v1.3-v2.0", "能力扩展与工程化", "自定义宠物持久化、串门、GitHub 陪学、安全加固"],
            ["2026.07.01", "情感捕捉细化", "reply/emotion/need/intensity/risk_level 五维结构"],
            ["2026.07.03", "Web 软件化", "三栏控制中心、状态与轻养成、桌宠预览与设置"],
            ["2026.07.06-07", "Electron 桌宠 MVP", "透明置顶、2 字气泡、轻聊天、托盘、会话共享、归属加固"],
        ],
        [1900, 2600, 4860],
        font_size=8.5,
        first_col_bold=True,
        alternate_fill=True,
    )

    add_subheading(doc, "三次关键取舍")
    add_bullet(doc, "先用规则做主动关怀：MVP 要先验证“主动是否被接受”，再把复杂触发交给 Agent。")
    add_bullet(doc, "保留原生 Web 前端：桌面化改造重点是窗口职责与交互形态，不为技术栈升级而重写。")
    add_bullet(doc, "气泡只做召回，不承载完整内容：把陪伴感、隐私和打扰成本放在同一个决策里权衡。")
    add_bullet(doc, "后台 Agent 不阻塞主回复：画像和情绪趋势允许失败降级，先守住核心对话可用性。")

    add_callout(doc, "迭代方法", "每个版本都同时检查用户可感知价值、技术边界和安全风险；规划文档只保留未完成项，已实现内容进入更新记录。")


def build_evidence(doc: Document) -> None:
    add_section_title(doc, "09  实现证据与个人贡献", "EVIDENCE & ROLE")
    add_subheading(doc, "仓库快照（2026-07-23）")
    add_table(
        doc,
        ["证据维度", "核验结果", "说明"],
        [
            ["核心代码规模", "15,040 行", "排除 node_modules 与 dist 后统计"],
            ["后端接口", "30 个 API 端点", "会话、聊天、自定义、串门、陪学"],
            ["数据模型", "12 张 SQLite 表", "关系、记忆、画像、串门、学习完整落库"],
            ["工程文件", "30 Python / 8 HTML / 9 CSS / 9 JS", "覆盖后端、Web 与 Electron"],
            ["测试资产", "5 个本地测试脚本", "记忆、学习归属、天气与 API 等"],
        ],
        [2100, 1900, 5360],
        font_size=8.8,
        first_col_bold=True,
        alternate_fill=True,
    )

    add_subheading(doc, "我的角色与产出")
    add_table(
        doc,
        ["能力", "具体产出"],
        [
            ["产品定义", "完成定位转型、用户场景、双终端职责、Phase 0-4 路线与验收标准"],
            ["AI 产品设计", "设计分层记忆、五维情感信号、后台画像 Agent、多角色上下文隔离"],
            ["交互与体验", "设计三栏控制中心、2 字气泡、轻聊天、勿扰与托盘路径"],
            ["原型实现", "推动 FastAPI、原生 Web、Electron、SQLite 形成可运行 MVP"],
            ["质量治理", "持续维护 bug / 安全 / 更新文档，完成身份归属、注入与 SSRF 等收口"],
        ],
        [2200, 7160],
        font_size=8.8,
        first_col_bold=True,
    )
    add_text(
        doc,
        "证据口径：以上为代码仓库可核验的实现规模，不等同于真实用户增长或商业结果；用户数据仍需上线验证。",
        size=8.8,
        color=MUTED,
        italic=True,
        after=0,
    )


def build_roadmap(doc: Document) -> None:
    add_section_title(doc, "10  路线图、验证指标与复盘", "NEXT & REFLECTION")
    add_subheading(doc, "下一阶段优先级")
    add_table(
        doc,
        ["优先级", "要做什么", "为什么现在做"],
        [
            ["P0 用户验证", "5-8 位目标用户连续使用 7 天，记录提醒接受度、记忆体验和留存原因", "先验证关系型桌宠是否产生真实复访"],
            ["P1 桌宠增强", "开机自启、位置记忆、多屏、通知隐私、应用数据目录", "把“能跑”提升到“适合长期挂着”"],
            ["P1 主动关怀", "统一主动服务、全局频控、天气穿衣建议与学习提醒", "减少重复逻辑并验证主动价值"],
            ["P2 工程质量", "全局 LLM 并发、事务加固、pytest + mock、结构化 tool calling", "为长期运行和迭代建立稳定底座"],
        ],
        [1500, 4850, 3010],
        font_size=8.5,
        first_col_bold=True,
        alternate_fill=True,
    )

    add_subheading(doc, "建议验证指标")
    add_table(
        doc,
        ["目标", "核心指标", "判断信号"],
        [
            ["关系是否成立", "D1 / D7 留存、每周主动会话天数", "用户不是只在新鲜期打开一次"],
            ["记忆是否有价值", "记忆命中满意度、错误记忆纠正率", "“记得我”多于“记错我”"],
            ["主动是否克制", "气泡点击率、关闭率、勿扰开启率", "触达带来回应，而不是逃离"],
            ["陪学是否可持续", "章节完成率、连续学习天数、宠物反馈互动率", "陪伴真正进入学习流程"],
        ],
        [2000, 3600, 3760],
        font_size=8.6,
        first_col_bold=True,
    )

    add_subheading(doc, "项目复盘")
    add_bullet(doc, "最重要的收获：AI 产品的差异化不只来自模型能力，而来自记忆、触发、人格、入口和安全边界的组合。")
    add_bullet(doc, "仍需改进：当前成果以工程验证为主，缺少真实用户研究、留存数据和成本看板；下一阶段应优先补齐证据链。")
    add_bullet(doc, "产品原则：让宠物更懂用户，但不越界；更主动，但不打扰；更拟人，但不冒充真人或专业服务。")

    add_spacer(doc, 8)
    add_text(doc, "谢谢阅读", size=18, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_text(doc, "QAgent Pet  |  王思勤", size=9.5, color=LIGHT_MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)


def build_portfolio() -> Document:
    doc = Document()
    configure_document(doc)

    pages = [
        build_cover,
        build_toc,
        build_overview,
        build_positioning,
        build_core_loop,
        build_ai_design,
        build_desktop,
        build_agents,
        build_architecture,
        build_timeline,
        build_evidence,
        build_roadmap,
    ]
    for index, builder in enumerate(pages):
        builder(doc)
        if index != len(pages) - 1:
            add_page_break(doc)

    return doc


def main() -> None:
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    document = build_portfolio()
    document.save(OUTPUT_DOCX)
    size_kb = os.path.getsize(OUTPUT_DOCX) / 1024
    print(f"Portfolio generated: {OUTPUT_DOCX}")
    print(f"Size: {size_kb:.1f} KB")


if __name__ == "__main__":
    main()
